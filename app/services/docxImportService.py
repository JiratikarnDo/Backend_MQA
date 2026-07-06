from io import BytesIO
from typing import Dict, List, Optional

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
import re


courseCodePattern = re.compile(r"(?:\d{2}|[Xx]{2})\s*[-–]\s*(?:\d{2}|[Xx]{2})\s*[-–]\s*(?:\d{3}|[Xx]{3})")
creditPattern = re.compile(r"\d+\s*\(\s*\d+\s*[-–]\s*\d+\s*[-–]\s*\d+\s*\)")
latinPattern = re.compile(r"[A-Za-z]")
thaiPattern = re.compile(r"[\u0E00-\u0E7F]")
thaiDigitMap = str.maketrans("๐๑๒๓๔๕๖๗๘๙", "0123456789")

metadataColumnKeywords = [
    "ความแตกต่าง",
    "หมายเหตุ",
    "สาระการปรับปรุง",
    "ประเด็น",
]


legacyThaiCharMap = str.maketrans({
    "\uf700": "่",
    "\uf701": "้",
    "\uf702": "๊",
    "\uf703": "๋",
    "\uf704": "์",
    "\uf705": "ํ",
    "\uf706": "ิ",
    "\uf707": "ี",
    "\uf708": "ึ",
    "\uf709": "ื",
    "\uf70a": "่",
    "\uf70b": "้",
    "\uf70c": "๊",
    "\uf70d": "๋",
    "\uf70e": "์",
    "\uf70f": "ํ",
    "\uf710": "ั",
    "\uf711": "ิ",
    "\uf712": "็",
    "\uf713": "ี",
    "\uf714": "ื",
    "\uf715": "ึ",
})


def normalizeLegacyThaiText(text: str) -> str:
    cleanText = str(text or "")

    # ไฟล์ Word เก่าหรือไฟล์ที่แปลงมาจาก .doc/PDF บางไฟล์ใช้ glyph legacy
    # เช่น เรียนรู, อาน, ปญหา ซึ่ง python-docx อ่านเป็น Unicode PUA
    # จึงแปลงกลับเป็นวรรณยุกต์/สระไทยก่อนนำไป parse
    cleanText = cleanText.translate(legacyThaiCharMap)

    legacyWordMap = {
        "กลุม": "กลุ่ม",
        "ใช": "ใช้",
        "ผู": "ผู้",
        "อาน": "อ่าน",
        "เกี่ยวของ": "เกี่ยวข้อง",
        "ออนไลน": "ออนไลน์",
        "จําเปน": "จำเป็น",
        "ปญหา": "ปัญหา",
        "กลยุทธ": "กลยุทธ์",
        "ธุรกิจสตารทอัพ": "ธุรกิจสตาร์ทอัพ",
        "ทํางาน": "ทำงาน",
        "นํา": "นำ",
        "สําหรับ": "สำหรับ",
        "สําเร็จ": "สำเร็จ",
        "สําคัญ": "สำคัญ",
    }

    for wrongText, correctText in legacyWordMap.items():
        cleanText = cleanText.replace(wrongText, correctText)

    # กรณีที่ Word แสดงเป็นกล่องจริง ๆ ให้ลบออกเพื่อไม่ให้ขวางการจับชื่อ/คำอธิบาย
    # ถ้าอักขระไทยเสียไปแล้วทั้งหมด โค้ดไม่สามารถกู้คำเดิมได้ แต่จะยัง import รายวิชาได้
    cleanText = re.sub(r"[□☐■�]", " ", cleanText)
    return cleanText


def normalizeText(text: str) -> str:
    cleanText = normalizeLegacyThaiText(text)
    cleanText = cleanText.replace("\xa0", " ")
    cleanText = cleanText.replace("\u200b", "")
    cleanText = re.sub(r"\s+", " ", cleanText)
    return cleanText.strip()


def normalizeYearText(text: str) -> str:
    return normalizeText(str(text or "").translate(thaiDigitMap))


def extractCurriculumYearsFromText(text: str) -> List[int]:
    cleanText = normalizeYearText(text)
    years: List[int] = []

    # ใช้เฉพาะปีที่ผูกกับคำว่า “หลักสูตร” เพื่อไม่ให้ไปหยิบปีงบประมาณ/ปีการศึกษา เช่น 2570-2573
    patterns = [
        r"หลักสูตร(?:เดิม|ปรับปรุง)?\s*\(?\s*พ\.?\s*ศ\.?\s*(\d{4})\s*\)?",
        r"หลักสูตร[^\n]{0,50}?พ\.?\s*ศ\.?\s*(\d{4})",
    ]

    for pattern in patterns:
        for match in re.finditer(pattern, cleanText):
            year = int(match.group(1))

            if 2400 <= year <= 2700 and year not in years:
                years.append(year)

    return years


def getLatestCurriculumYearFromDocument(document) -> Optional[int]:
    texts = []

    for paragraph in document.paragraphs:
        if paragraph.text:
            texts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    texts.append(cell.text)

    years = extractCurriculumYearsFromText("\n".join(texts))

    if not years:
        return None

    return max(years)


def isMetadataColumn(text: str) -> bool:
    cleanText = normalizeYearText(text)
    return any(keyword in cleanText for keyword in metadataColumnKeywords)


def getTableCurriculumYears(table) -> List[int]:
    tableText = "\n".join(
        cell.text
        for row in table.rows
        for cell in row.cells
    )

    return extractCurriculumYearsFromText(tableText)


def isCurriculumComparisonTable(table, latestCurriculumYear: Optional[int]) -> bool:
    years = getTableCurriculumYears(table)

    if len(years) >= 2:
        return True

    if not latestCurriculumYear:
        return False

    tableText = normalizeYearText(
        " ".join(
            cell.text
            for row in table.rows
            for cell in row.cells
        )
    )

    return str(latestCurriculumYear) in tableText and "ความแตกต่าง" in tableText


def shouldSkipTableBecauseNotLatest(table, latestCurriculumYear: Optional[int]) -> bool:
    if not latestCurriculumYear:
        return False

    years = getTableCurriculumYears(table)

    if not years:
        return False

    tableText = normalizeYearText(
        " ".join(
            cell.text
            for row in table.rows
            for cell in row.cells
        )
    )

    # ข้ามเฉพาะตารางหลักสูตรที่มีปีหลักสูตร แต่ไม่มีปีล่าสุด
    if "หลักสูตร" in tableText and latestCurriculumYear not in years:
        return True

    return False


def getLatestCurriculumColumnIndexes(
    table,
    latestCurriculumYear: Optional[int],
) -> List[int]:
    if not latestCurriculumYear:
        return []

    latestYearText = str(latestCurriculumYear)

    for row in table.rows[:12]:
        cellTexts = [normalizeYearText(cell.text) for cell in row.cells]
        latestIndexes = []

        for index, cellText in enumerate(cellTexts):
            if latestYearText not in cellText:
                continue

            if "หลักสูตร" not in cellText and "ปรับปรุง" not in cellText:
                continue

            latestIndexes.append(index)

        if latestIndexes:
            resultIndexes = set(latestIndexes)

            # หาคอลัมน์หน่วยกิตที่อยู่ถัดจากคอลัมน์หลักสูตรล่าสุด
            for index in latestIndexes:
                nextIndex = index + 1

                if nextIndex >= len(cellTexts):
                    continue

                nextCellText = cellTexts[nextIndex]

                if isMetadataColumn(nextCellText):
                    continue

                # กรณีหัวตาราง merge: คอลัมน์ถัดไปอาจยังเป็นหลักสูตรล่าสุด หรือเป็นหน่วยกิต
                if latestYearText in nextCellText or "หน่วยกิต" in nextCellText or nextCellText == "":
                    resultIndexes.add(nextIndex)
                    continue

                # กรณีทั่วไป เช่น [รหัสวิชา, ชื่อวิชา, หน่วยกิต] ฝั่งล่าสุด
                if "หลักสูตร" not in nextCellText and "ความแตกต่าง" not in nextCellText:
                    resultIndexes.add(nextIndex)

            # เพิ่มอีก 1 ช่องหลังกลุ่มล่าสุด ถ้าเป็นหน่วยกิต
            maxIndex = max(resultIndexes)
            nextIndex = maxIndex + 1

            if nextIndex < len(cellTexts):
                nextCellText = cellTexts[nextIndex]

                if "หน่วยกิต" in nextCellText and not isMetadataColumn(nextCellText):
                    resultIndexes.add(nextIndex)

            return sorted(resultIndexes)

    if isCurriculumComparisonTable(table, latestCurriculumYear):
        columnCount = max(len(row.cells) for row in table.rows)

        # โครงสร้างที่พบบ่อยในไฟล์นี้:
        # 2564... | หน่วยกิต | 2569... | หน่วยกิต | ความแตกต่าง
        if columnCount >= 7:
            return list(range(columnCount - 4, columnCount - 1))

        if columnCount >= 5:
            return list(range(columnCount - 3, columnCount - 1))

        halfIndex = columnCount // 2
        return list(range(halfIndex, columnCount))

    return []


def getRowLinesForLatestCurriculum(
    row,
    table,
    latestCurriculumYear: Optional[int],
) -> List[str]:
    if shouldSkipTableBecauseNotLatest(table, latestCurriculumYear):
        return []

    latestColumnIndexes = getLatestCurriculumColumnIndexes(table, latestCurriculumYear)

    if latestColumnIndexes:
        lines: List[str] = []

        for index in latestColumnIndexes:
            if index >= len(row.cells):
                continue

            for line in splitCellLines(row.cells[index].text):
                if line and line not in lines and not isMetadataColumn(line):
                    lines.append(line)

        return lines

    return getUniqueRowLines(row)


def splitCellLines(text: str) -> List[str]:
    lines = []

    for line in str(text or "").splitlines():
        cleanLine = normalizeText(line)
        if cleanLine:
            lines.append(cleanLine)

    return lines


def normalizeCourseCode(courseCode: str) -> str:
    cleanCode = normalizeText(courseCode)
    cleanCode = cleanCode.replace("–", "-")
    cleanCode = re.sub(r"\s*[-]\s*", "-", cleanCode)
    return cleanCode.upper() if cleanCode.lower().startswith("xx-") else cleanCode


def normalizeCreditText(creditText: str) -> str:
    return normalizeText(creditText).replace("–", "-").replace(" ", "")


def hasThaiText(text: str) -> bool:
    return bool(thaiPattern.search(text or ""))


def getUniqueRowLines(row) -> List[str]:
    lines = []

    for cell in row.cells:
        for line in splitCellLines(cell.text):
            if line and line not in lines:
                lines.append(line)

    return lines


def iterBlockItems(document):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def extractSubjectCodesFromText(text: str) -> List[str]:
    codes = []

    for match in courseCodePattern.finditer(text or ""):
        courseCode = normalizeCourseCode(match.group(0))
        if courseCode not in codes:
            codes.append(courseCode)

    return codes


def parseCreditText(creditText: Optional[str]) -> Dict[str, int]:
    match = re.search(
        r"(\d+)\s*\(\s*(\d+)\s*[-–]\s*(\d+)\s*[-–]\s*(\d+)\s*\)",
        str(creditText or ""),
    )

    if not match:
        return {
            "totalCredits": 0,
            "lectureHours": 0,
            "labHours": 0,
            "selfStudyHours": 0,
        }

    return {
        "totalCredits": int(match.group(1)),
        "lectureHours": int(match.group(2)),
        "labHours": int(match.group(3)),
        "selfStudyHours": int(match.group(4)),
    }


def parseCourseName(nameText: str) -> Dict[str, str]:
    lines = splitCellLines(nameText)
    cleanLines = []

    skipWords = [
        "ให้ศึกษารายวิชาต่อไปนี้",
        "ให้เลือกศึกษารายวิชาต่อไปนี้",
        "และให้เลือกศึกษาจากกลุ่มรายวิชาต่อไปนี้",
        "เลือกรายวิชาสหกิจศึกษา",
        "หรือรายวิชาฝึกประสบการณ์วิชาชีพ",
        "หรือ",
    ]

    for line in lines:
        cleanLine = courseCodePattern.sub("", line)
        cleanLine = creditPattern.sub("", cleanLine)
        cleanLine = normalizeText(cleanLine)

        if not cleanLine:
            continue

        if any(skipWord in cleanLine for skipWord in skipWords):
            continue

        cleanLines.append(cleanLine)

    if not cleanLines:
        return {"courseNameThai": "", "courseNameEnglish": ""}

    englishStartIndex = None

    for index, line in enumerate(cleanLines):
        if latinPattern.search(line):
            englishStartIndex = index
            break

    if englishStartIndex is None:
        return {
            "courseNameThai": normalizeText(" ".join(cleanLines)),
            "courseNameEnglish": "",
        }

    return {
        "courseNameThai": normalizeText(" ".join(cleanLines[:englishStartIndex])),
        "courseNameEnglish": normalizeText(" ".join(cleanLines[englishStartIndex:])),
    }


def resolveSubjectCategoryFromCourseCode(courseCode: str) -> Dict[str, str]:
    cleanCourseCode = normalizeCourseCode(courseCode)

    if cleanCourseCode.lower().startswith("xx-"):
        return {"subjectCategory": "freeElective", "subCategory": "วิชาเลือกเสรี"}

    if cleanCourseCode.startswith("15-"):
        return {
            "subjectCategory": "generalEducation",
            "subCategory": "หมวดวิชาศึกษาทั่วไป",
        }

    if cleanCourseCode.startswith("04-"):
        return {"subjectCategory": "specific", "subCategory": "หมวดวิชาเฉพาะ"}

    return {"subjectCategory": "specific", "subCategory": "หมวดวิชาเฉพาะ"}


def resolveCategoryContext(
    text: str,
    currentSubjectCategory: str,
    currentSubCategory: str,
) -> Dict[str, str]:
    cleanText = normalizeText(text)
    subjectCategory = currentSubjectCategory
    subCategory = currentSubCategory

    if re.search(r"(?:^|\s)1\.?\s*หมวดวิชาศึกษาทั่วไป", cleanText) or "หมวดรายวิชาศึกษาทั่วไป" in cleanText:
        subjectCategory = "generalEducation"
        subCategory = "หมวดวิชาศึกษาทั่วไป"

    elif re.search(r"(?:^|\s)2\.?\s*หมวดวิชาเฉพาะ", cleanText) or "หมวดวิชาเฉพาะ" in cleanText:
        subjectCategory = "specific"
        subCategory = "หมวดวิชาเฉพาะ"

    elif (re.search(r"(?:^|\s)3\.?\s*หมวด", cleanText) and "เลือกเสรี" in cleanText) or "หมวดวิชาเลือกเสรี" in cleanText or "หมวดเลือกเสรี" in cleanText:
        subjectCategory = "freeElective"
        subCategory = "วิชาเลือกเสรี"

    generalGroupPatterns = [
        (r"กลุ่มสาระวิชาอัตลักษณ์|กลุ่มวิชาอัตลักษณ์", "กลุ่มสาระวิชาอัตลักษณ์"),
        (r"กลุ่มสาระวิชาคุณภาพชีวิต", "กลุ่มสาระวิชาคุณภาพชีวิต"),
        (r"กลุ่มสาระวิชาคุณภาพการทำงาน", "กลุ่มสาระวิชาคุณภาพการทำงาน"),
        (r"กลุ่มสาระวิชาภาษาและการสื่อสาร", "กลุ่มสาระวิชาภาษาและการสื่อสาร"),
        (r"กลุ่มสาระวิชาการปรับตัวและการใช้ชีวิต", "กลุ่มสาระวิชาการปรับตัวและการใช้ชีวิต"),
        (r"กลุ่มสาระวิชาความเป็นพลเมืองไทยและพลเมืองโลก", "กลุ่มสาระวิชาความเป็นพลเมืองไทยและพลเมืองโลก"),
    ]

    specificGroupPatterns = [
        (r"2\.1\s*กลุ่มวิชาแกน|กลุ่มวิชาแกน", "กลุ่มวิชาแกน"),
        (r"2\.2\s*กลุ่มวิชาบังคับ|กลุ่มวิชาบังคับ", "กลุ่มวิชาบังคับ"),
        (r"2\.3\s*กลุ่มวิชาเลือก", "กลุ่มวิชาเลือก"),
        (r"กลุ่มวิชาการพัฒนาซอฟต์แวร์|กลุ่มการพัฒนาซอฟต์แวร์", "กลุ่มวิชาการพัฒนาซอฟต์แวร์"),
        (r"กลุ่มวิทยาการข้อมูล", "กลุ่มวิทยาการข้อมูล"),
        (r"กลุ่มปัญญาประดิษฐ์และระบบอัตโนมัติ", "กลุ่มปัญญาประดิษฐ์และระบบอัตโนมัติ"),
        (r"กลุ่มโครงสร้างพื้นฐานทางเทคโนโลยีสารสนเทศ", "กลุ่มโครงสร้างพื้นฐานทางเทคโนโลยีสารสนเทศ"),
        (r"2\.4\s*กลุ่มวิชาเสริมสร้างประสบการณ์|กลุ่มวิชาเสริมสร้างประสบการณ์ทางวิทยาการคอมพิวเตอร์|กลุ่มวิชาเสริมสร้างประสบการณ์วิชาชีพ", "กลุ่มวิชาเสริมสร้างประสบการณ์วิชาชีพ"),
    ]

    for pattern, groupName in generalGroupPatterns:
        if re.search(pattern, cleanText):
            subjectCategory = "generalEducation"
            subCategory = groupName
            break

    for pattern, groupName in specificGroupPatterns:
        if re.search(pattern, cleanText):
            subjectCategory = "specific"
            subCategory = groupName
            break

    return {"subjectCategory": subjectCategory, "subCategory": subCategory}


def resolveCourseCategoryForParsedCourse(
    courseCode: str,
    currentSubjectCategory: str,
    currentSubCategory: str,
) -> Dict[str, str]:
    cleanCourseCode = normalizeCourseCode(courseCode)

    if cleanCourseCode.lower().startswith("xx-"):
        return {"subjectCategory": "freeElective", "subCategory": "วิชาเลือกเสรี"}

    if cleanCourseCode.startswith("15-"):
        return {
            "subjectCategory": "generalEducation",
            "subCategory": currentSubCategory
            if "กลุ่ม" in currentSubCategory and currentSubjectCategory == "generalEducation"
            else "หมวดวิชาศึกษาทั่วไป",
        }

    if currentSubjectCategory == "freeElective":
        return {"subjectCategory": "freeElective", "subCategory": "วิชาเลือกเสรี"}

    if cleanCourseCode.startswith("04-"):
        return {
            "subjectCategory": "specific",
            "subCategory": currentSubCategory if currentSubjectCategory == "specific" else "หมวดวิชาเฉพาะ",
        }

    if currentSubjectCategory:
        return {"subjectCategory": currentSubjectCategory, "subCategory": currentSubCategory}

    return resolveSubjectCategoryFromCourseCode(cleanCourseCode)


def isCategoryHeadingOnly(text: str) -> bool:
    cleanText = normalizeText(text)

    if not cleanText:
        return False

    if courseCodePattern.search(cleanText):
        return False

    categoryKeywords = [
        "หมวดวิชาศึกษาทั่วไป",
        "หมวดรายวิชาศึกษาทั่วไป",
        "หมวดวิชาเฉพาะ",
        "หมวดวิชาเลือกเสรี",
        "หมวดเลือกเสรี",
        "กลุ่มสาระวิชา",
        "กลุ่มวิชา",
    ]

    return any(keyword in cleanText for keyword in categoryKeywords)


def extractEnglishCourseNameAfterCredit(text: str) -> str:
    lines = splitCellLines(text)

    for line in lines:
        cleanLine = normalizeText(line)
        cleanLine = courseCodePattern.sub("", cleanLine)
        cleanLine = creditPattern.sub("", cleanLine)
        cleanLine = normalizeText(cleanLine)

        if not cleanLine:
            continue

        if hasThaiText(cleanLine):
            continue

        if not latinPattern.search(cleanLine):
            continue

        if courseCodePattern.search(cleanLine):
            continue

        # กันไม่ให้เอาคำอธิบายภาษาอังกฤษยาว ๆ มาเป็นชื่อวิชา
        wordCount = len(cleanLine.split())

        if wordCount > 8:
            continue

        return cleanLine

    return ""


def parseCoursesFromTextBlock(
    textBlock: str,
    subjectCategory: str = "specific",
    subCategory: str = "หมวดวิชาเฉพาะ",
) -> List[Dict]:
    courses = []
    matches = list(courseCodePattern.finditer(textBlock or ""))

    if not matches:
        return courses

    for index, match in enumerate(matches):
        courseCode = normalizeCourseCode(match.group(0))

        if courseCode.lower().startswith("xx-"):
            continue

        startIndex = match.end()
        endIndex = matches[index + 1].start() if index + 1 < len(matches) else len(textBlock)
        courseChunk = textBlock[startIndex:endIndex]
        creditMatch = creditPattern.search(courseChunk)

        if not creditMatch:
            continue

        creditText = normalizeCreditText(creditMatch.group(0))
        beforeCreditText = courseChunk[: creditMatch.start()]
        afterCreditText = courseChunk[creditMatch.end():]

        parsedName = parseCourseName(beforeCreditText)
        englishNameAfterCredit = extractEnglishCourseNameAfterCredit(afterCreditText)

        # รองรับไฟล์ Word ที่วางชื่ออังกฤษไว้บรรทัดถัดจากหน่วยกิต
        # เช่น 15-03-005 ผู้ประกอบการนวัตกรรม 3(2-2-5) แล้วบรรทัดถัดไปเป็น Innovative Entrepreneur
        if not parsedName["courseNameEnglish"] and englishNameAfterCredit:
            parsedName["courseNameEnglish"] = englishNameAfterCredit

        if not parsedName["courseNameThai"] and not parsedName["courseNameEnglish"]:
            continue

        creditValue = parseCreditText(creditText)
        categoryValue = resolveCourseCategoryForParsedCourse(courseCode, subjectCategory, subCategory)

        courses.append(
            {
                "courseCode": courseCode,
                "curriculumLevel": "bachelor",
                "courseNameThai": parsedName["courseNameThai"],
                "courseNameEnglish": parsedName["courseNameEnglish"],
                "subjectCategory": categoryValue["subjectCategory"],
                "subCategory": categoryValue["subCategory"],
                "studyLine": "สายวิทยาศาสตร์",
                "totalCredits": creditValue["totalCredits"],
                "lectureHours": creditValue["lectureHours"],
                "labHours": creditValue["labHours"],
                "selfStudyHours": creditValue["selfStudyHours"],
                "descriptionThai": "",
                "descriptionEnglish": "",
                "hasPreSubjects": "no",
                "preSubjects": [],
                "hasCoSubjects": "no",
                "coSubjects": [],
            }
        )

    return courses



def collectDocumentLinesInOrder(document) -> List[str]:
    lines = []

    for block in iterBlockItems(document):
        if isinstance(block, Paragraph):
            paragraphText = normalizeText(block.text)

            if paragraphText:
                lines.append(paragraphText)

        elif isinstance(block, Table):
            for row in block.rows:
                rowLines = getUniqueRowLines(row)

                for line in rowLines:
                    if line:
                        lines.append(line)

    return lines


def lineHasOnlyOneCourseCode(text: str) -> bool:
    return len(extractSubjectCodesFromText(text)) == 1


def isLikelyCourseHeaderStart(lines: List[str], index: int) -> bool:
    if index >= len(lines):
        return False

    currentLine = normalizeText(lines[index])

    # บรรทัดความสัมพันธ์ เช่น "วิชาบังคับก่อน: 04-10-102 ..."
    # มีรหัสวิชาอยู่ก็จริง แต่ไม่ใช่หัวรายวิชาใหม่
    relationHeaderKeywords = [
        "วิชาที่ต้องเรียนมาก่อน",
        "วิชาบังคับก่อน",
        "วิชาที่ต้องเรียนพร้อมกัน",
        "วิชาบังคับร่วม",
        "prerequisite",
        "pre-requisite",
        "corequisite",
        "co-requisite",
    ]

    if any(keyword in currentLine.lower() for keyword in relationHeaderKeywords):
        return False

    if not lineHasOnlyOneCourseCode(currentLine):
        return False

    courseCode = extractSubjectCodesFromText(currentLine)[0]

    if courseCode.lower().startswith("xx-"):
        return False

    # ถ้ารหัสวิชาอยู่กลางประโยคยาว ๆ มักเป็นอ้างอิง/วิชาบังคับก่อน ไม่ใช่หัวข้อรายวิชา
    textWithoutCode = removeCourseCodeAndCredit(currentLine)
    if len(textWithoutCode) > 80:
        return False

    lookAheadEnd = min(index + 7, len(lines))

    for lookAheadIndex in range(index, lookAheadEnd):
        lookAheadLine = normalizeText(lines[lookAheadIndex])

        if any(keyword in lookAheadLine.lower() for keyword in relationHeaderKeywords):
            continue

        if creditPattern.search(lookAheadLine or ""):
            return True

    return False


def removeCourseCodeAndCredit(text: str) -> str:
    cleanText = normalizeText(text)
    cleanText = courseCodePattern.sub(" ", cleanText)
    cleanText = creditPattern.sub(" ", cleanText)
    cleanText = normalizeText(cleanText)
    return cleanText


def isLikelyEnglishCourseNameLine(text: str) -> bool:
    cleanText = normalizeText(text)

    if not cleanText:
        return False

    if hasThaiText(cleanText):
        return False

    if not latinPattern.search(cleanText):
        return False

    if courseCodePattern.search(cleanText) or creditPattern.search(cleanText):
        return False

    # ชื่อวิชาภาษาอังกฤษมักสั้นกว่าคำอธิบาย
    return len(cleanText) <= 120


def findNextCourseHeaderIndex(lines: List[str], startIndex: int) -> int:
    for index in range(startIndex, len(lines)):
        if isLikelyCourseHeaderStart(lines, index):
            return index

    return len(lines)


def extractPrerequisiteAndCorequisite(text: str) -> Dict[str, List[str]]:
    preSubjects = []
    coSubjects = []

    prePatterns = [
        r"วิชาที่ต้องเรียนมาก่อน\s*[:：]?\s*([^\n]+)",
        r"วิชาบังคับก่อน\s*[:：]?\s*([^\n]+)",
        r"Prerequisite\s*[:：]?\s*([^\n]+)",
        r"Pre-requisite\s*[:：]?\s*([^\n]+)",
    ]

    coPatterns = [
        r"วิชาที่ต้องเรียนพร้อมกัน\s*[:：]?\s*([^\n]+)",
        r"วิชาบังคับร่วม\s*[:：]?\s*([^\n]+)",
        r"Corequisite\s*[:：]?\s*([^\n]+)",
        r"Co-requisite\s*[:：]?\s*([^\n]+)",
    ]

    for pattern in prePatterns:
        match = re.search(pattern, text or "", re.IGNORECASE)

        if match:
            preSubjects.extend(extractSubjectCodesFromText(match.group(1)))

    for pattern in coPatterns:
        match = re.search(pattern, text or "", re.IGNORECASE)

        if match:
            coSubjects.extend(extractSubjectCodesFromText(match.group(1)))

    return {
        "preSubjects": list(dict.fromkeys(preSubjects)),
        "coSubjects": list(dict.fromkeys(coSubjects)),
    }


def isRelationLine(text: str) -> bool:
    cleanText = normalizeText(text).lower()

    relationKeywords = [
        "วิชาที่ต้องเรียนมาก่อน",
        "วิชาบังคับก่อน",
        "prerequisite",
        "pre-requisite",
        "วิชาที่ต้องเรียนพร้อมกัน",
        "วิชาบังคับร่วม",
        "corequisite",
        "co-requisite",
    ]

    return any(keyword in cleanText for keyword in relationKeywords)


def isEnglishDescriptionStartLine(text: str) -> bool:
    cleanText = normalizeText(text)

    if not cleanText:
        return False

    if hasThaiText(cleanText):
        return False

    if not latinPattern.search(cleanText):
        return False

    if courseCodePattern.search(cleanText) or creditPattern.search(cleanText):
        return False

    return True


def cleanDescriptionLines(lines: List[str]) -> List[str]:
    cleanLines = []

    for line in lines:
        cleanLine = removeCourseCodeAndCredit(line)

        if not cleanLine:
            continue

        if isRelationLine(cleanLine):
            continue

        blockedKeywords = [
            "หมวดวิชาเฉพาะ",
            "หมวดวิชาศึกษาทั่วไป",
            "คำอธิบายรายวิชา",
            "หลักสูตรวิทยาศาสตรบัณฑิต",
            "สาขาวิชาวิทยาการคอมพิวเตอร์",
            "กลุ่มวิชาแกน",
            "กลุ่มวิชาบังคับ",
            "กลุ่มวิชาเลือก",
            "ภาคผนวก",
        ]

        if any(keyword in cleanLine for keyword in blockedKeywords):
            continue

        cleanLines.append(cleanLine)

    return cleanLines


def splitDescriptionByLanguage(lines: List[str]) -> Dict[str, str]:
    cleanLines = cleanDescriptionLines(lines)

    if not cleanLines:
        return {
            "descriptionThai": "",
            "descriptionEnglish": "",
        }

    thaiLines = []
    englishLines = []
    isEnglishPart = False

    for line in cleanLines:
        cleanLine = normalizeText(line)

        if not cleanLine:
            continue

        if isEnglishDescriptionStartLine(cleanLine):
            isEnglishPart = True

        if isEnglishPart:
            englishLines.append(cleanLine)
        else:
            thaiLines.append(cleanLine)

    return {
        "descriptionThai": normalizeText(" ".join(thaiLines)),
        "descriptionEnglish": normalizeText(" ".join(englishLines)),
    }


def parseCourseDescriptionHeader(lines: List[str], index: int) -> Optional[Dict]:
    if not isLikelyCourseHeaderStart(lines, index):
        return None

    courseCode = extractSubjectCodesFromText(lines[index])[0]

    if courseCode.lower().startswith("xx-"):
        return None

    nextHeaderIndex = findNextCourseHeaderIndex(lines, index + 1)
    segmentEnd = min(nextHeaderIndex, index + 20, len(lines))
    segment = lines[index:segmentEnd]

    creditIndex = None

    for offset, line in enumerate(segment):
        if creditPattern.search(line or ""):
            creditIndex = offset
            break

    if creditIndex is None:
        return None

    headerLines = []

    firstLineName = removeCourseCodeAndCredit(segment[0])

    if firstLineName:
        headerLines.append(firstLineName)

    for offset in range(1, creditIndex):
        candidateLine = removeCourseCodeAndCredit(segment[offset])

        if candidateLine:
            headerLines.append(candidateLine)

    headerEndOffset = creditIndex + 1

    # บางไฟล์วางชื่ออังกฤษหลังหน่วยกิต เช่น 04-10-101 / ระบบปฏิบัติการ / 3(2-2-5) / Operating Systems
    if headerEndOffset < len(segment) and isLikelyEnglishCourseNameLine(segment[headerEndOffset]):
        headerLines.append(segment[headerEndOffset])
        headerEndOffset += 1

    parsedName = parseCourseName("\n".join(headerLines))

    if not parsedName["courseNameThai"] and not parsedName["courseNameEnglish"]:
        return None

    return {
        "courseCode": courseCode,
        "courseNameThai": parsedName["courseNameThai"],
        "courseNameEnglish": parsedName["courseNameEnglish"],
        "descriptionStartIndex": index + headerEndOffset,
        "nextHeaderIndex": nextHeaderIndex,
    }


def normalizeNameForMatch(text: str) -> str:
    cleanText = normalizeText(text)
    cleanText = courseCodePattern.sub("", cleanText)
    cleanText = creditPattern.sub("", cleanText)
    cleanText = re.sub(r"\s+", "", cleanText)
    return cleanText.lower()


def isSameCourseName(sourceCourse: Dict, candidateCourse: Dict) -> bool:
    sourceNameThai = normalizeNameForMatch(sourceCourse.get("courseNameThai") or "")
    sourceNameEnglish = normalizeNameForMatch(sourceCourse.get("courseNameEnglish") or "")
    candidateNameThai = normalizeNameForMatch(candidateCourse.get("courseNameThai") or "")
    candidateNameEnglish = normalizeNameForMatch(candidateCourse.get("courseNameEnglish") or "")

    thaiMatched = sourceNameThai and candidateNameThai and sourceNameThai == candidateNameThai
    englishMatched = sourceNameEnglish and candidateNameEnglish and sourceNameEnglish == candidateNameEnglish

    return bool(thaiMatched or englishMatched)


def extractDescriptionSectionLines(document) -> List[str]:
    lines = collectDocumentLinesInOrder(document)
    startIndex = None
    endIndex = None

    for index, line in enumerate(lines):
        cleanLine = normalizeText(line)

        # เลี่ยงสารบัญช่วงต้นเล่ม โดยเริ่มจากภาคผนวกจริงช่วงท้ายเอกสาร
        if index < 1000:
            continue

        if (
            "คำอธิบายรายวิชาหมวดศึกษาทั่วไป" in cleanLine
            or "คำอธิบายรายวิชาหมวดวิชาเฉพาะ" in cleanLine
            or ("ภาคผนวก 7" in cleanLine and "คำอธิบายรายวิชา" in cleanLine)
        ):
            startIndex = index
            break

    if startIndex is None:
        return []

    for index in range(startIndex + 1, len(lines)):
        cleanLine = normalizeText(lines[index])

        if (
            "ภาคผนวก 9" in cleanLine
            or "ภาคผนวกที่ 9" in cleanLine
            or "เอกสารความร่วมมือ" in cleanLine
        ):
            endIndex = index
            break

    if endIndex is None:
        endIndex = len(lines)

    return lines[startIndex:endIndex]


def extractCourseDescriptionMapFromDocument(document) -> Dict[str, List[Dict]]:
    lines = extractDescriptionSectionLines(document)
    descriptionMap: Dict[str, List[Dict]] = {}
    index = 0

    while index < len(lines):
        headerData = parseCourseDescriptionHeader(lines, index)

        if not headerData:
            index += 1
            continue

        descriptionStartIndex = headerData["descriptionStartIndex"]
        nextHeaderIndex = headerData["nextHeaderIndex"]
        descriptionLines = lines[descriptionStartIndex:nextHeaderIndex]
        descriptionParts = splitDescriptionByLanguage(descriptionLines)
        descriptionTextLength = len(descriptionParts["descriptionThai"]) + len(descriptionParts["descriptionEnglish"])

        if descriptionTextLength >= 20:
            relationData = extractPrerequisiteAndCorequisite("\n".join(descriptionLines))
            courseCode = headerData["courseCode"]
            candidate = {
                "courseCode": courseCode,
                "courseNameThai": headerData.get("courseNameThai") or "",
                "courseNameEnglish": headerData.get("courseNameEnglish") or "",
                "descriptionThai": descriptionParts["descriptionThai"],
                "descriptionEnglish": descriptionParts["descriptionEnglish"],
                "preSubjects": relationData["preSubjects"],
                "coSubjects": relationData["coSubjects"],
            }
            descriptionMap.setdefault(courseCode, []).append(candidate)

        index += 1

    return descriptionMap


def findMatchedDescriptionData(parsedCourse: Dict, descriptionCandidates: List[Dict]) -> Dict:
    for candidate in descriptionCandidates:
        if isSameCourseName(parsedCourse, candidate):
            return candidate

    if descriptionCandidates:
        return descriptionCandidates[0]

    return {
        "descriptionThai": "",
        "descriptionEnglish": "",
        "preSubjects": [],
        "coSubjects": [],
    }


def extractCoursesFromDocx(fileBytes: bytes) -> List[Dict]:
    document = Document(BytesIO(fileBytes))
    latestCurriculumYear = getLatestCurriculumYearFromDocument(document)
    descriptionMap = extractCourseDescriptionMapFromDocument(document)
    courses = []
    seenCourseCodes = set()
    currentSubjectCategory = "specific"
    currentSubCategory = "หมวดวิชาเฉพาะ"

    print("LATEST CURRICULUM YEAR:", latestCurriculumYear)

    for block in iterBlockItems(document):
        if isinstance(block, Paragraph):
            paragraphText = normalizeText(block.text)

            if not paragraphText:
                continue

            contextValue = resolveCategoryContext(
                paragraphText,
                currentSubjectCategory,
                currentSubCategory,
            )
            currentSubjectCategory = contextValue["subjectCategory"]
            currentSubCategory = contextValue["subCategory"]
            continue

        if not isinstance(block, Table):
            continue

        if shouldSkipTableBecauseNotLatest(block, latestCurriculumYear):
            continue

        latestColumnIndexes = getLatestCurriculumColumnIndexes(block, latestCurriculumYear)

        for row in block.rows:
            if latestColumnIndexes:
                rowLines = []

                for columnIndex in latestColumnIndexes:
                    if columnIndex >= len(row.cells):
                        continue

                    for line in splitCellLines(row.cells[columnIndex].text):
                        if line and line not in rowLines and not isMetadataColumn(line):
                            rowLines.append(line)
            else:
                rowLines = getUniqueRowLines(row)

            if not rowLines:
                continue

            rowTextBlock = "\n".join(rowLines)
            contextValue = resolveCategoryContext(
                rowTextBlock,
                currentSubjectCategory,
                currentSubCategory,
            )
            currentSubjectCategory = contextValue["subjectCategory"]
            currentSubCategory = contextValue["subCategory"]

            if isCategoryHeadingOnly(rowTextBlock):
                continue

            parsedCourses = parseCoursesFromTextBlock(
                rowTextBlock,
                currentSubjectCategory,
                currentSubCategory,
            )

            for parsedCourse in parsedCourses:
                courseCode = normalizeCourseCode(parsedCourse["courseCode"])

                if courseCode in seenCourseCodes:
                    continue

                descriptionData = findMatchedDescriptionData(
                    parsedCourse,
                    descriptionMap.get(courseCode, []),
                )

                parsedCourse["descriptionThai"] = descriptionData["descriptionThai"]
                parsedCourse["descriptionEnglish"] = descriptionData["descriptionEnglish"]
                parsedCourse["preSubjects"] = descriptionData["preSubjects"]
                parsedCourse["coSubjects"] = descriptionData["coSubjects"]
                parsedCourse["hasPreSubjects"] = "yes" if descriptionData["preSubjects"] else "no"
                parsedCourse["hasCoSubjects"] = "yes" if descriptionData["coSubjects"] else "no"

                seenCourseCodes.add(courseCode)
                courses.append(parsedCourse)

    print("DOCX IMPORT SUMMARY:", {
        "total": len(courses),
        "generalEducation": len([c for c in courses if c.get("subjectCategory") == "generalEducation"]),
        "specific": len([c for c in courses if c.get("subjectCategory") == "specific"]),
        "freeElective": len([c for c in courses if c.get("subjectCategory") == "freeElective"]),
    })

    return courses
