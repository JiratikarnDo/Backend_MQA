from io import BytesIO
from pathlib import Path
import os
import shutil
import subprocess
import tempfile
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from fastapi import HTTPException


THAI_FONT_NAME = "TH Sarabun New"


def findSofficePath() -> str:
    sofficePath = shutil.which("soffice") or shutil.which("libreoffice")

    if sofficePath:
        return sofficePath

    windowsPaths = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]

    for path in windowsPaths:
        if Path(path).exists():
            return path

    return ""


def setRunFont(run, fontName: str = THAI_FONT_NAME) -> None:
    run.font.name = fontName

    runProperties = run._element.get_or_add_rPr()
    runFonts = runProperties.rFonts

    if runFonts is None:
        runFonts = runProperties._add_rFonts()

    runFonts.set(qn("w:ascii"), fontName)
    runFonts.set(qn("w:hAnsi"), fontName)
    runFonts.set(qn("w:eastAsia"), fontName)
    runFonts.set(qn("w:cs"), fontName)


def applyThaiFontToParagraphs(paragraphs, fontName: str = THAI_FONT_NAME) -> None:
    for paragraph in paragraphs:
        for run in paragraph.runs:
            setRunFont(run, fontName)


def applyThaiFontToTables(tables, fontName: str = THAI_FONT_NAME) -> None:
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                applyThaiFontToParagraphs(cell.paragraphs, fontName)
                applyThaiFontToTables(cell.tables, fontName)


def applyThaiFontToDocument(document: Document, fontName: str = THAI_FONT_NAME) -> None:
    # ตั้ง style เท่าที่ python-docx เขียนได้ เพื่อให้ Word/LibreOffice ใช้ฟอนต์ไทยนี้ตอน render
    for style in document.styles:
        try:
            if hasattr(style, "font") and style.font:
                style.font.name = fontName

            styleElement = style.element
            runProperties = styleElement.get_or_add_rPr()
            runFonts = runProperties.rFonts

            if runFonts is None:
                runFonts = runProperties._add_rFonts()

            runFonts.set(qn("w:ascii"), fontName)
            runFonts.set(qn("w:hAnsi"), fontName)
            runFonts.set(qn("w:eastAsia"), fontName)
            runFonts.set(qn("w:cs"), fontName)
        except Exception:
            # บาง style เป็น built-in/locked ถ้าแก้ไม่ได้ให้ข้าม ไม่ให้ระบบ import ล่ม
            continue

    applyThaiFontToParagraphs(document.paragraphs, fontName)
    applyThaiFontToTables(document.tables, fontName)


def applyThaiFontToDocxBytes(fileBytes: bytes, fontName: str = THAI_FONT_NAME) -> bytes:
    try:
        document = Document(BytesIO(fileBytes))
    except Exception as error:
        raise HTTPException(
            status_code=400,
            detail=f"อ่านไฟล์ .docx ไม่สำเร็จ: {str(error)}",
        )

    applyThaiFontToDocument(document, fontName)

    outputBuffer = BytesIO()
    document.save(outputBuffer)
    return outputBuffer.getvalue()


def convertDocToDocxBytes(fileBytes: bytes, fileName: str) -> bytes:
    sofficePath = findSofficePath()

    if not sofficePath:
        raise HTTPException(
            status_code=500,
            detail=(
                "ไม่พบ LibreOffice ในเครื่อง Backend จึงแปลงไฟล์ .doc ไม่ได้ "
                "กรุณาติดตั้ง LibreOffice หรือแปลงไฟล์เป็น .docx ก่อนอัปโหลด"
            ),
        )

    safeFileName = Path(fileName or "input.doc").name

    with tempfile.TemporaryDirectory() as tempDir:
        tempPath = Path(tempDir)
        inputPath = tempPath / safeFileName
        outputDir = tempPath / "converted"
        userInstallDir = tempPath / "lo-profile"
        outputDir.mkdir(parents=True, exist_ok=True)
        userInstallDir.mkdir(parents=True, exist_ok=True)

        inputPath.write_bytes(fileBytes)

        env = os.environ.copy()
        env["SAL_USE_VCLPLUGIN"] = "svp"

        command = [
            sofficePath,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--nodefault",
            "--nolockcheck",
            f"-env:UserInstallation=file:///{str(userInstallDir).replace(os.sep, '/')}",
            "--convert-to",
            "docx",
            "--outdir",
            str(outputDir),
            str(inputPath),
        ]

        try:
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                timeout=120,
                env=env,
            )
        except subprocess.TimeoutExpired:
            raise HTTPException(
                status_code=500,
                detail="แปลงไฟล์ .doc ไม่สำเร็จ: LibreOffice ใช้เวลานานเกินไป",
            )

        outputFiles = list(outputDir.glob("*.docx"))

        if result.returncode != 0 or not outputFiles:
            raise HTTPException(
                status_code=500,
                detail=(
                    "แปลงไฟล์ .doc เป็น .docx ไม่สำเร็จ "
                    f"stdout={result.stdout} stderr={result.stderr}"
                ),
            )

        convertedBytes = outputFiles[0].read_bytes()

    return applyThaiFontToDocxBytes(convertedBytes)


def convertWordFileToDocxBytes(fileBytes: bytes, fileName: Optional[str]) -> bytes:
    suffix = Path(fileName or "").suffix.lower()

    if suffix == ".docx":
        return applyThaiFontToDocxBytes(fileBytes)

    if suffix == ".doc":
        return convertDocToDocxBytes(fileBytes, fileName or "input.doc")

    raise HTTPException(
        status_code=400,
        detail="รองรับเฉพาะไฟล์ .doc และ .docx เท่านั้น",
    )
